import re
from io import BytesIO
from typing import List, Optional

import pdfplumber
import docx
from pypdf import PdfReader

# -------------------------
# Text extraction functions
# -------------------------
def extract_text_from_pdf(file_obj) -> str:
    # pdfplumber accepts a file-like object, but make sure we use BytesIO for django UploadedFile
    if not hasattr(file_obj, "read"):
        # if a path string passed
        with pdfplumber.open(file_obj) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)

    data = file_obj.read()
    # keep a BytesIO copy so callers can still access file if needed
    buf = BytesIO(data)
    with pdfplumber.open(buf) as pdf:
        text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    # ensure file pointer position reset for Django (optional)
    try:
        file_obj.seek(0)
    except Exception:
        pass
    return text


def extract_text_from_docx(file_obj) -> str:
    """
    Accepts file-like object (UploadedFile). python-docx needs a file path OR file-like with .read
    """
    # python-docx can accept a file-like object that supports seek()
    if hasattr(file_obj, "read"):
        data = file_obj.read()
        buf = BytesIO(data)
        doc = docx.Document(buf)
        try:
            file_obj.seek(0)
        except Exception:
            pass
    else:
        doc = docx.Document(file_obj)

    return "\n".join(p.text for p in doc.paragraphs)


def extract_text_from_txt(file_obj) -> str:
    if hasattr(file_obj, "read"):
        raw = file_obj.read()
        if isinstance(raw, bytes):
            try:
                return raw.decode("utf-8", errors="ignore")
            finally:
                try:
                    file_obj.seek(0)
                except Exception:
                    pass
        return raw
    else:
        with open(file_obj, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def extract_text(file_obj, filename: str) -> str:
    """
    Convenience wrapper - detect file type by extension and extract text.
    filename should be the uploaded file name (file.name).
    """
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_obj)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file_obj)
    elif name_lower.endswith(".txt"):
        return extract_text_from_txt(file_obj)
    else:
        # fallback: try pdf first, then docx, then raw decode
        try:
            return extract_text_from_pdf(file_obj)
        except Exception:
            try:
                return extract_text_from_docx(file_obj)
            except Exception:
                return extract_text_from_txt(file_obj)


# -------------------------
# Regex / heuristic extractors
# -------------------------
EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Za-z]{2,}", re.I)
# phone: capture +country or numbers with separators; keep digits and leading +
PHONE_REGEX = re.compile(r"(\+?\d{1,3}[\s\-\.]?\(?\d{1,4}\)?[\s\-\.]?\d{3,4}[\s\-\.]?\d{3,4})")
EDUCATION_KEYWORDS = [
    "bachelor", "master", "b\.tech", "m\.tech", "bsc", "msc", "mba", "phd", "bs", "ms", "b\.e", "m\.e", "bca", "mca"
]


def extract_emails(text: str) -> List[str]:
    return list({m.group(0).strip() for m in EMAIL_REGEX.finditer(text)})


def extract_phones(text: str) -> List[str]:
    matches = [m.group(0) for m in PHONE_REGEX.finditer(text)]
    cleaned = []
    for p in matches:
        p_clean = re.sub(r"[^\d+]", "", p)
        if len(re.sub(r"\D", "", p_clean)) >= 7:
            cleaned.append(p_clean)
    seen = set()
    out = []
    for p in cleaned:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out


def extract_education(text: str):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    edu = []

    edu_keywords = [
        "bachelor", "master", "bca", "mca", "b.sc", "m.sc",
        "b.tech", "m.tech", "b.e", "m.e", "phd", "12th", "10th"
    ]

    for ln in lines:
        lower = ln.lower()
        if any(kw in lower for kw in edu_keywords):
            # ignore lines containing URLs or bullets
            if not ln.startswith("●") and "http" not in ln:
                edu.append(ln)

    return edu



def guess_name(text: str):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    # Look at first 5–10 lines (resume header)
    header = lines[:10]

    # Try combining 2 consecutive lines for name
    for i in range(len(header) - 1):
        line1 = header[i]
        line2 = header[i + 1]

        # Skip if email or phone in these lines
        if EMAIL_REGEX.search(line1) or EMAIL_REGEX.search(line2):
            continue

        # If both lines contain mostly letters → likely a first/last name
        if line1.replace(" ", "").isalpha() and line2.replace(" ", "").isalpha():
            name = f"{line1} {line2}".strip()
            if 3 <= len(name) <= 40:
                return name.title()

    # fallback: your original logic
    for ln in header:
        if EMAIL_REGEX.search(ln) or PHONE_REGEX.search(ln):
            continue
        cleaned = re.sub(r"[^A-Za-z\s]", "", ln)
        if len(cleaned.strip()) > 1 and len(cleaned.split()) <= 4:
            return cleaned.title()

    return None



# -------------------------
# Single entrypoint
# -------------------------
def parse_resume(file_obj, filename: str) -> dict:
    text = extract_text(file_obj, filename) or ""
    emails = extract_emails(text)
    phones = extract_phones(text)
    education = extract_education(text)
    name = guess_name(text)

    return {
        "name": name,
        "emails": emails,
        "phones": phones,
        "education": education,
        "raw_text": text,
    }
